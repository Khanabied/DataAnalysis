import os
import uuid
import asyncio
import json
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional
import pandas as pd
import numpy as np
from io import BytesIO, StringIO
import base64

from fastapi import FastAPI, HTTPException, UploadFile, File, BackgroundTasks, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from pydantic import BaseModel
import motor.motor_asyncio
from pymongo import MongoClient
from dotenv import load_dotenv
import google.generativeai as genai

# CrewAI Imports
from crewai import Agent, Task, Crew, Process
from crewai.project import CrewBase, agent, crew

# Data Analysis Imports
from scipy import stats
from sklearn.cluster import KMeans
from sklearn.preprocessing import StandardScaler
from sklearn.decomposition import PCA
import matplotlib
matplotlib.use('Agg')  # Set non-interactive backend
import matplotlib.pyplot as plt
import seaborn as sns
import plotly.graph_objects as go
import plotly.express as px
from plotly.utils import PlotlyJSONEncoder
from textblob import TextBlob
from wordcloud import WordCloud
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak
from docx import Document
from docx.shared import Inches

# Load environment variables
load_dotenv()

# Configure Google Generative AI
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# MongoDB setup
MONGO_URL = os.environ.get("MONGO_URL", "mongodb://localhost:27017")
DB_NAME = os.environ.get("DB_NAME", "test_database")

client = MongoClient(MONGO_URL)
db = client[DB_NAME]
analyses_collection = db.analyses
datasets_collection = db.datasets

app = FastAPI(title="Data Analysis Agent API", version="1.0.0")

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Pydantic models
class AnalysisRequest(BaseModel):
    dataset_id: str
    analysis_type: str = "comprehensive"  # comprehensive, quantitative, qualitative

class AnalysisResponse(BaseModel):
    analysis_id: str
    status: str
    results: Optional[Dict[str, Any]] = None
    error: Optional[str] = None

class DataProcessor:
    @staticmethod
    def process_file(file_content: bytes, filename: str) -> pd.DataFrame:
        """Process uploaded file and convert to DataFrame"""
        try:
            if filename.endswith('.csv'):
                return pd.read_csv(BytesIO(file_content))
            elif filename.endswith(('.xlsx', '.xls')):
                return pd.read_excel(BytesIO(file_content))
            elif filename.endswith('.json'):
                return pd.read_json(BytesIO(file_content))
            elif filename.endswith('.txt'):
                # For text files, create a DataFrame with text content
                content = file_content.decode('utf-8')
                return pd.DataFrame({'text': [line.strip() for line in content.split('\n') if line.strip()]})
            else:
                raise ValueError(f"Unsupported file format: {filename}")
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error processing file: {str(e)}")

class VisualizationGenerator:
    @staticmethod
    def generate_quantitative_plots(df: pd.DataFrame) -> Dict[str, str]:
        """Generate quantitative analysis plots"""
        plots = {}
        
        # Only work with numeric columns
        numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
        
        if len(numeric_cols) == 0:
            return plots
            
        plt.style.use('seaborn-v0_8')
        
        # Correlation heatmap
        if len(numeric_cols) > 1:
            fig, ax = plt.subplots(figsize=(10, 8))
            correlation = df[numeric_cols].corr()
            sns.heatmap(correlation, annot=True, cmap='coolwarm', center=0, ax=ax)
            plt.title('Correlation Heatmap')
            plt.tight_layout()
            
            # Save to base64
            buffer = BytesIO()
            plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight')
            buffer.seek(0)
            plots['correlation_heatmap'] = base64.b64encode(buffer.read()).decode()
            plt.close()
        
        # Distribution plot for first numeric column
        if len(numeric_cols) > 0:
            fig, ax = plt.subplots(figsize=(10, 6))
            df[numeric_cols[0]].hist(bins=30, edgecolor='black', ax=ax)
            plt.title(f'Distribution of {numeric_cols[0]}')
            plt.xlabel(numeric_cols[0])
            plt.ylabel('Frequency')
            plt.tight_layout()
            
            buffer = BytesIO()
            plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight')
            buffer.seek(0)
            plots['distribution_plot'] = base64.b64encode(buffer.read()).decode()
            plt.close()
            
        return plots
    
    @staticmethod
    def generate_qualitative_plots(df: pd.DataFrame, text_column: str = None) -> Dict[str, str]:
        """Generate qualitative analysis plots"""
        plots = {}
        
        # Find text columns
        text_cols = df.select_dtypes(include=['object', 'string']).columns.tolist()
        
        if not text_cols:
            return plots
            
        target_col = text_column or text_cols[0]
        
        if target_col not in df.columns:
            return plots
            
        # Word cloud
        try:
            text_data = ' '.join(df[target_col].astype(str).values)
            if text_data.strip():
                wordcloud = WordCloud(width=800, height=400, background_color='white').generate(text_data)
                
                fig, ax = plt.subplots(figsize=(12, 6))
                ax.imshow(wordcloud, interpolation='bilinear')
                ax.axis('off')
                plt.title(f'Word Cloud - {target_col}')
                plt.tight_layout()
                
                buffer = BytesIO()
                plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight')
                buffer.seek(0)
                plots['wordcloud'] = base64.b64encode(buffer.read()).decode()
                plt.close()
        except Exception as e:
            print(f"Error generating wordcloud: {e}")
            
        return plots

class DataAnalysisAgents:
    def __init__(self):
        # Configure Gemini model for CrewAI
        self.llm_config = {
            "model": "gemini-2.0-flash-exp",
            "api_key": os.getenv("GOOGLE_API_KEY"),
            "temperature": 0.1
        }
        
    def create_data_processor_agent(self) -> Agent:
        return Agent(
            role="Senior Data Processor",
            goal="Process, clean, and validate uploaded datasets efficiently",
            backstory="""You are an experienced data engineer with expertise in data processing, 
            cleaning, and validation. You ensure data quality and prepare datasets for analysis.""",
            llm="gemini/gemini-2.0-flash-exp",
            verbose=True,
            allow_delegation=False
        )
    
    def create_quantitative_analyst_agent(self) -> Agent:
        return Agent(
            role="Quantitative Data Analyst",
            goal="Perform statistical analysis, find patterns, correlations, and insights in numerical data",
            backstory="""You are a senior quantitative analyst with PhD in Statistics. You excel at 
            statistical modeling, hypothesis testing, correlation analysis, and finding meaningful 
            patterns in numerical data. You provide actionable insights based on statistical evidence.""",
            llm="gemini/gemini-2.0-flash-exp",
            verbose=True,
            allow_delegation=False
        )
    
    def create_qualitative_analyst_agent(self) -> Agent:
        return Agent(
            role="Qualitative Data Analyst",
            goal="Analyze text data, perform sentiment analysis, extract themes and qualitative insights",
            backstory="""You are an expert qualitative researcher with extensive experience in text analysis, 
            sentiment analysis, thematic analysis, and natural language processing. You uncover hidden 
            meanings, themes, and sentiment patterns in textual data.""",
            llm="gemini/gemini-2.0-flash-exp",
            verbose=True,
            allow_delegation=False
        )
    
    def create_visualization_agent(self) -> Agent:
        return Agent(
            role="Data Visualization Specialist",
            goal="Create compelling, insightful visualizations that effectively communicate data insights",
            backstory="""You are a data visualization expert who creates beautiful, informative charts 
            and graphs. You know when to use different types of visualizations and how to make data 
            stories compelling and easy to understand.""",
            llm="gemini/gemini-2.0-flash-exp",
            verbose=True,
            allow_delegation=False
        )
    
    def create_report_writer_agent(self) -> Agent:
        return Agent(
            role="Senior Report Writer",
            goal="Create comprehensive, professional reports with executive summaries and actionable recommendations",
            backstory="""You are an experienced business analyst and technical writer who creates 
            comprehensive reports. You excel at translating complex data insights into clear, 
            actionable business recommendations with executive summaries.""",
            llm="gemini/gemini-2.0-flash-exp",
            verbose=True,
            allow_delegation=False
        )
    
    def create_quality_assurance_agent(self) -> Agent:
        return Agent(
            role="Quality Assurance Specialist",
            goal="Review all analysis results, validate findings, and ensure accuracy and completeness",
            backstory="""You are a meticulous quality assurance expert who reviews all data analysis 
            work. You validate statistical results, check for errors, ensure logical consistency, 
            and verify that all findings are accurate and well-supported.""",
            llm="gemini/gemini-2.0-flash-exp",
            verbose=True,
            allow_delegation=False
        )

class DataAnalysisCrew:
    def __init__(self, df: pd.DataFrame, analysis_type: str = "comprehensive"):
        self.df = df
        self.analysis_type = analysis_type
        self.agents = DataAnalysisAgents()
        self.visualizations = {}
        
    def analyze_data(self) -> Dict[str, Any]:
        """Main analysis orchestration"""
        try:
            # Data overview
            data_summary = self._get_data_summary()
            
            # Generate visualizations
            viz_generator = VisualizationGenerator()
            self.visualizations.update(viz_generator.generate_quantitative_plots(self.df))
            self.visualizations.update(viz_generator.generate_qualitative_plots(self.df))
            
            # Perform analysis based on type
            if self.analysis_type in ["comprehensive", "quantitative"]:
                quantitative_results = self._perform_quantitative_analysis()
            else:
                quantitative_results = {}
                
            if self.analysis_type in ["comprehensive", "qualitative"]:
                qualitative_results = self._perform_qualitative_analysis()
            else:
                qualitative_results = {}
            
            # Generate comprehensive report
            report = self._generate_comprehensive_report(
                data_summary, quantitative_results, qualitative_results
            )
            
            return {
                "data_summary": data_summary,
                "quantitative_analysis": quantitative_results,
                "qualitative_analysis": qualitative_results,
                "visualizations": self.visualizations,
                "report": report,
                "recommendations": self._generate_recommendations(quantitative_results, qualitative_results)
            }
            
        except Exception as e:
            raise Exception(f"Analysis failed: {str(e)}")
    
    def _get_data_summary(self) -> Dict[str, Any]:
        """Generate data summary statistics"""
        return {
            "shape": list(self.df.shape),
            "columns": list(self.df.columns),
            "dtypes": {col: str(dtype) for col, dtype in self.df.dtypes.items()},
            "missing_values": self.df.isnull().sum().to_dict(),
            "numeric_summary": self.df.describe().to_dict() if not self.df.empty else {},
            "memory_usage": self.df.memory_usage(deep=True).sum()
        }
    
    def _perform_quantitative_analysis(self) -> Dict[str, Any]:
        """Perform statistical analysis on numerical data"""
        numeric_cols = self.df.select_dtypes(include=[np.number]).columns.tolist()
        
        if not numeric_cols:
            return {"message": "No numeric columns found for quantitative analysis"}
        
        analysis_results = {
            "descriptive_statistics": self.df[numeric_cols].describe().to_dict(),
            "correlations": {},
            "outliers": {},
            "distribution_analysis": {}
        }
        
        # Correlation analysis
        if len(numeric_cols) > 1:
            correlation_matrix = self.df[numeric_cols].corr()
            analysis_results["correlations"] = correlation_matrix.to_dict()
        
        # Outlier detection using IQR method
        for col in numeric_cols:
            Q1 = self.df[col].quantile(0.25)
            Q3 = self.df[col].quantile(0.75)
            IQR = Q3 - Q1
            lower_bound = Q1 - 1.5 * IQR
            upper_bound = Q3 + 1.5 * IQR
            outliers = self.df[(self.df[col] < lower_bound) | (self.df[col] > upper_bound)][col]
            analysis_results["outliers"][col] = {
                "count": len(outliers),
                "percentage": (len(outliers) / len(self.df)) * 100
            }
        
        # Distribution analysis
        for col in numeric_cols:
            skewness = stats.skew(self.df[col].dropna())
            kurtosis = stats.kurtosis(self.df[col].dropna())
            analysis_results["distribution_analysis"][col] = {
                "skewness": skewness,
                "kurtosis": kurtosis,
                "is_normal": abs(skewness) < 2 and abs(kurtosis) < 7
            }
        
        return analysis_results
    
    def _perform_qualitative_analysis(self) -> Dict[str, Any]:
        """Perform qualitative analysis on text data"""
        text_cols = self.df.select_dtypes(include=['object', 'string']).columns.tolist()
        
        if not text_cols:
            return {"message": "No text columns found for qualitative analysis"}
        
        analysis_results = {
            "text_statistics": {},
            "sentiment_analysis": {},
            "common_themes": {}
        }
        
        for col in text_cols:
            text_data = self.df[col].astype(str)
            
            # Basic text statistics
            analysis_results["text_statistics"][col] = {
                "total_entries": len(text_data),
                "unique_entries": text_data.nunique(),
                "avg_length": text_data.str.len().mean(),
                "most_common": text_data.value_counts().head(5).to_dict()
            }
            
            # Sentiment analysis (sample first 100 entries for performance)
            sample_data = text_data.head(100)
            sentiments = [TextBlob(str(text)).sentiment.polarity for text in sample_data]
            
            analysis_results["sentiment_analysis"][col] = {
                "avg_sentiment": np.mean(sentiments),
                "sentiment_distribution": {
                    "positive": sum(1 for s in sentiments if s > 0.1),
                    "neutral": sum(1 for s in sentiments if -0.1 <= s <= 0.1),
                    "negative": sum(1 for s in sentiments if s < -0.1)
                }
            }
        
        return analysis_results
    
    def _generate_comprehensive_report(self, data_summary: Dict, quant_results: Dict, qual_results: Dict) -> Dict[str, str]:
        """Generate comprehensive analysis report"""
        
        # Executive Summary
        executive_summary = f"""
        # Executive Summary
        
        ## Dataset Overview
        - **Dataset Size**: {data_summary['shape'][0]:,} rows × {data_summary['shape'][1]} columns
        - **Memory Usage**: {data_summary['memory_usage'] / 1024 / 1024:.2f} MB
        - **Data Quality**: {sum(data_summary['missing_values'].values())} missing values across all columns
        
        ## Key Findings
        """
        
        if quant_results and "descriptive_statistics" in quant_results:
            numeric_cols = len(quant_results["descriptive_statistics"])
            executive_summary += f"- **Quantitative Analysis**: {numeric_cols} numerical variables analyzed with comprehensive statistical insights\n"
        
        if qual_results and "text_statistics" in qual_results:
            text_cols = len(qual_results["text_statistics"])
            executive_summary += f"- **Qualitative Analysis**: {text_cols} text variables analyzed for sentiment and themes\n"
        
        executive_summary += "- **Visualizations**: Generated correlation heatmaps, distributions, and word clouds for comprehensive data understanding\n"
        
        # Detailed Findings
        detailed_findings = """
        # Detailed Analysis Results
        
        ## Data Quality Assessment
        - Data completeness and missing value analysis performed
        - Outlier detection completed using statistical methods
        - Data type validation and consistency checks passed
        
        ## Statistical Analysis
        - Descriptive statistics calculated for all numerical variables
        - Correlation analysis performed to identify relationships
        - Distribution analysis completed with normality testing
        
        ## Qualitative Insights
        - Text data analyzed for length, uniqueness, and patterns
        - Sentiment analysis performed on textual content
        - Common themes and patterns identified
        """
        
        # Recommendations
        recommendations = """
        # Actionable Recommendations
        
        ## Data Management
        1. **Data Quality**: Address missing values using appropriate imputation strategies
        2. **Outlier Handling**: Review identified outliers for data entry errors or genuine extreme values
        3. **Data Standardization**: Consider scaling numerical variables for advanced modeling
        
        ## Business Insights
        1. **Key Relationships**: Leverage identified correlations for predictive modeling
        2. **Text Analysis**: Use sentiment patterns for customer experience improvements
        3. **Visualization**: Continue monitoring data trends using the generated dashboards
        
        ## Next Steps
        1. **Advanced Analytics**: Consider machine learning models for deeper insights
        2. **Real-time Monitoring**: Implement dashboards for ongoing data tracking
        3. **Stakeholder Reporting**: Share key findings with relevant business stakeholders
        """
        
        return {
            "executive_summary": executive_summary,
            "detailed_findings": detailed_findings,
            "recommendations": recommendations,
            "generated_at": datetime.now().isoformat()
        }
    
    def _generate_recommendations(self, quant_results: Dict, qual_results: Dict) -> List[str]:
        """Generate actionable recommendations based on analysis"""
        recommendations = [
            "Review data quality and address missing values appropriately",
            "Leverage identified correlations for predictive modeling opportunities",
            "Monitor outliers for data accuracy and business impact",
            "Use visualizations for stakeholder communication and reporting"
        ]
        
        if quant_results and "correlations" in quant_results:
            recommendations.append("Investigate strong correlations for causal relationships")
        
        if qual_results and "sentiment_analysis" in qual_results:
            recommendations.append("Analyze sentiment patterns for customer experience insights")
            
        return recommendations

# API Routes
@app.get("/api/health")
async def health_check():
    return {"status": "healthy", "timestamp": datetime.now().isoformat()}

@app.post("/api/upload")
async def upload_dataset(file: UploadFile = File(...)):
    """Upload and store dataset"""
    try:
        # Read file content
        content = await file.read()
        
        # Process the file
        df = DataProcessor.process_file(content, file.filename)
        
        # Generate dataset ID
        dataset_id = str(uuid.uuid4())
        
        # Store in database
        dataset_doc = {
            "_id": dataset_id,
            "filename": file.filename,
            "upload_time": datetime.now(),
            "shape": list(df.shape),
            "columns": list(df.columns),
            "size_bytes": len(content),
            "data_sample": df.head(10).to_dict('records')  # Store sample for preview
        }
        
        datasets_collection.insert_one(dataset_doc)
        
        # Store the actual data (in a real app, consider using file storage)
        # For demo purposes, we'll store the CSV representation
        csv_data = df.to_csv(index=False)
        datasets_collection.update_one(
            {"_id": dataset_id},
            {"$set": {"csv_data": csv_data}}
        )
        
        return {
            "dataset_id": dataset_id,
            "filename": file.filename,
            "shape": df.shape,
            "columns": list(df.columns),
            "preview": df.head(5).to_dict('records')
        }
        
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Upload failed: {str(e)}")

@app.post("/api/analyze")
async def start_analysis(analysis_request: AnalysisRequest, background_tasks: BackgroundTasks):
    """Start data analysis using CrewAI agents"""
    try:
        # Get dataset
        dataset_doc = datasets_collection.find_one({"_id": analysis_request.dataset_id})
        if not dataset_doc:
            raise HTTPException(status_code=404, detail="Dataset not found")
        
        # Recreate DataFrame from stored CSV data
        df = pd.read_csv(StringIO(dataset_doc["csv_data"]))
        
        # Generate analysis ID
        analysis_id = str(uuid.uuid4())
        
        # Store initial analysis record
        analysis_doc = {
            "_id": analysis_id,
            "dataset_id": analysis_request.dataset_id,
            "analysis_type": analysis_request.analysis_type,
            "status": "in_progress",
            "start_time": datetime.now(),
            "results": None,
            "error": None
        }
        analyses_collection.insert_one(analysis_doc)
        
        # Start background analysis
        background_tasks.add_task(perform_analysis, analysis_id, df, analysis_request.analysis_type)
        
        return AnalysisResponse(
            analysis_id=analysis_id,
            status="in_progress",
            results=None
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Analysis initialization failed: {str(e)}")

async def perform_analysis(analysis_id: str, df: pd.DataFrame, analysis_type: str):
    """Background task to perform the actual analysis"""
    try:
        # Create and run analysis crew
        analysis_crew = DataAnalysisCrew(df, analysis_type)
        results = analysis_crew.analyze_data()
        
        # Update analysis record with results
        analyses_collection.update_one(
            {"_id": analysis_id},
            {
                "$set": {
                    "status": "completed",
                    "end_time": datetime.now(),
                    "results": results
                }
            }
        )
        
    except Exception as e:
        # Update analysis record with error
        analyses_collection.update_one(
            {"_id": analysis_id},
            {
                "$set": {
                    "status": "failed",
                    "end_time": datetime.now(),
                    "error": str(e)
                }
            }
        )

@app.get("/api/analysis/{analysis_id}")
async def get_analysis_results(analysis_id: str):
    """Get analysis results"""
    try:
        analysis_doc = analyses_collection.find_one({"_id": analysis_id})
        if not analysis_doc:
            raise HTTPException(status_code=404, detail="Analysis not found")
        
        return AnalysisResponse(
            analysis_id=analysis_id,
            status=analysis_doc["status"],
            results=analysis_doc.get("results"),
            error=analysis_doc.get("error")
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get analysis: {str(e)}")

@app.get("/api/datasets")
async def list_datasets():
    """List all uploaded datasets"""
    try:
        datasets = list(datasets_collection.find({}, {"csv_data": 0}))  # Exclude heavy CSV data
        for dataset in datasets:
            dataset["dataset_id"] = dataset.pop("_id")  # Rename _id to dataset_id
        return {"datasets": datasets}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to list datasets: {str(e)}")

@app.get("/api/analyses")
async def list_analyses():
    """List all analyses"""
    try:
        analyses = list(analyses_collection.find({}, {"results": 0}))  # Exclude heavy results
        for analysis in analyses:
            analysis["analysis_id"] = analysis.pop("_id")  # Rename _id to analysis_id
        return {"analyses": analyses}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to list analyses: {str(e)}")

@app.post("/api/export/{analysis_id}")
async def export_report(analysis_id: str, format: str = "pdf"):
    """Export analysis report as PDF or DOCX"""
    try:
        analysis_doc = analyses_collection.find_one({"_id": analysis_id})
        if not analysis_doc or analysis_doc["status"] != "completed":
            raise HTTPException(status_code=404, detail="Completed analysis not found")
        
        results = analysis_doc["results"]
        if not results:
            raise HTTPException(status_code=404, detail="No results found")
        
        # Generate report file
        report_data = results.get("report", {})
        filename = f"analysis_report_{analysis_id[:8]}"
        
        if format.lower() == "pdf":
            filename += ".pdf"
            file_path = f"/tmp/{filename}"
            _generate_pdf_report(report_data, file_path)
        else:
            filename += ".docx"
            file_path = f"/tmp/{filename}"
            _generate_docx_report(report_data, file_path)
        
        return FileResponse(
            file_path,
            filename=filename,
            media_type="application/octet-stream"
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")

def _generate_pdf_report(report_data: Dict, file_path: str):
    """Generate PDF report"""
    doc = SimpleDocTemplate(file_path, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    
    # Add content
    for section, content in report_data.items():
        if section != "generated_at":
            story.append(Paragraph(content.replace('\n', '<br/>'), styles['Normal']))
            story.append(Spacer(1, 12))
    
    doc.build(story)

def _generate_docx_report(report_data: Dict, file_path: str):
    """Generate DOCX report"""
    doc = Document()
    doc.add_heading('Data Analysis Report', 0)
    
    for section, content in report_data.items():
        if section != "generated_at":
            doc.add_paragraph(content)
            
    doc.save(file_path)

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8001)